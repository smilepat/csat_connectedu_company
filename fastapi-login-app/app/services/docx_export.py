# app/services/export_docx.py
import re, os, base64, time
from io import BytesIO
from tempfile import NamedTemporaryFile
from typing import Optional, List, Any, Tuple

from app.schemas.export_docx import (
    ExportPayload, ExportItem, SubItem, LabeledOption,
    ImageSpec, ChartData, TableData, ChartDataset
)

# ── 지연 임포트: 무거운 의존성은 함수 내부에서 import ──────────────

FONT_KO_CANDIDATES = ["NanumGothic", "Noto Sans CJK KR", "Noto Sans CJK KR Regular", "Apple SD Gothic Neo", "Malgun Gothic"]
FONT_SYMBOL_CANDIDATES = ["Noto Sans Symbols2", "Segoe UI Symbol"]
NBSP = "\u00A0"

def _pick_font(candidates):
    try:
        import matplotlib
        matplotlib.use("Agg")
        from matplotlib import font_manager as fm
        names = {f.name for f in fm.fontManager.ttflist}
        for name in candidates:
            if name in names:
                return name
    except Exception:
        pass
    return candidates[0]

FONT_KO = _pick_font(FONT_KO_CANDIDATES)
FONT_SYMBOL = _pick_font(FONT_SYMBOL_CANDIDATES)

# ── 유틸 ───────────────────────────────────────────────────────────
def _strip_html(s: Optional[str]) -> str:
    if not s:
        return ""
    s = re.sub(r"(?i)<br\s*/?>", "\n", s)
    s = re.sub(r"<[^>]+>", "", s)
    return re.sub(r"\n{3,}", "\n\n", s).strip()

def _strip_controls(s: str) -> str:
    return re.sub(r"[\u200B-\u200D\uFEFF]", "", s)

def circled_label(i: int) -> str:
    base = 0x2460
    return chr(base + i) if 0 <= i <= 19 else f"{i+1}."

def _data_url_to_bytes(data_url: str) -> bytes:
    m = re.match(r"^data:(image/[\w\+\-\.]+);base64,(.+)$", data_url or "")
    if not m:
        if re.match(r"^[A-Za-z0-9+/=\s]+$", data_url or ""):
            return base64.b64decode(data_url)
        raise ValueError("Invalid data URL")
    return base64.b64decode(m.group(2))

def _strip_html_except_u(s: Optional[str]) -> str:
    if not s:
        return ""
    s = re.sub(r"(?i)<br\s*/?>", "\n", s)
    s = re.sub(r"(?i)</u>", "__CLOSE_U__", s)
    s = re.sub(r"(?i)<u>", "__OPEN_U__", s)
    s = re.sub(r"<[^>]+>", "", s)
    return s.replace("__OPEN_U__", "<u>").replace("__CLOSE_U__", "</u>")

def _preserve_spaces(text: str) -> str:
    text = text.replace("\t", "    ")
    return re.sub(r" {2,}", lambda m: NBSP * len(m.group(0)), text)

# ── <u>/<br> 지원 런 작성기 ────────────────────────────────────────
def add_rich_ko(par, html_text: Optional[str]):
    if not html_text:
        return
    raw = _strip_controls(_strip_html_except_u(html_text))
    if not raw:
        return
    lines = raw.replace("\r\n", "\n").replace("\r", "\n").split("\n")
    for li, line in enumerate(lines):
        if li > 0:
            par.add_run().add_break()
        parts = re.split(r"(?i)(<u>.*?</u>)", line)
        for part in parts:
            if not part:
                continue
            m = re.match(r"(?i)^<u>(.*?)</u>$", part)
            if m:
                txt = _preserve_spaces(m.group(1))
                run = add_ko_run(par, txt)
                run.underline = True
            else:
                txt = _preserve_spaces(part)
                add_ko_run(par, txt)

# ── DOCX 빌드 보조 ─────────────────────────────────────────────────
def _docx_primitives():
    from docx import Document
    from docx.text.paragraph import Paragraph
    from docx.text.run import Run
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    from docx.shared import Mm
    from docx.enum.table import WD_TABLE_ALIGNMENT
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    return Document, Paragraph, Run, qn, OxmlElement, Mm, WD_TABLE_ALIGNMENT, WD_ALIGN_PARAGRAPH

def _plt():
    import matplotlib
    matplotlib.use("Agg")
    import matplotlib.pyplot as plt
    os.environ.setdefault("MPLCONFIGDIR", "/tmp/mplconfig")
    try:
        os.makedirs("/tmp/mplconfig", exist_ok=True)
    except Exception:
        pass
    try:
        plt.rcParams["font.family"] = FONT_KO
    except Exception:
        pass
    return plt

def add_ko_run(par, text: str):
    _, _, _, qn, _, _, _, _ = _docx_primitives()
    run = par.add_run(text)
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.get_or_add_rFonts()
    run.font.name = FONT_KO
    rFonts.set(qn("w:ascii"),    FONT_KO)
    rFonts.set(qn("w:hAnsi"),    FONT_KO)
    rFonts.set(qn("w:eastAsia"), FONT_KO)
    rFonts.set(qn("w:cs"),       FONT_KO)
    return run

def add_symbol_run(par, text: str):
    _, _, _, qn, _, _, _, _ = _docx_primitives()
    run = par.add_run(text)
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.get_or_add_rFonts()
    run.font.name = FONT_SYMBOL
    rFonts.set(qn("w:ascii"),    FONT_SYMBOL)
    rFonts.set(qn("w:hAnsi"),    FONT_SYMBOL)
    rFonts.set(qn("w:eastAsia"), FONT_SYMBOL)
    rFonts.set(qn("w:cs"),       FONT_SYMBOL)
    return run

def set_cell_borders(cell, top=("single", 12, "000000"),
                          left=("single", 12, "000000"),
                          bottom=("single", 12, "000000"),
                          right=("single", 12, "000000")):
    _, _, _, qn, OxmlElement, _, _, _ = _docx_primitives()
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = tcPr.find(qn("w:tcBorders"))
    if tcBorders is None:
        tcBorders = OxmlElement("w:tcBorders")
        tcPr.append(tcBorders)

    def _edge(tag, spec):
        if spec is None:
            return
        val, sz, color = spec
        edge = tcBorders.find(qn(f"w:{tag}"))
        if edge is None:
            edge = OxmlElement(f"w:{tag}")
            tcBorders.append(edge)
        edge.set(qn("w:val"), val)
        edge.set(qn("w:sz"), str(sz))
        edge.set(qn("w:color"), color)

    _edge("top", top); _edge("left", left); _edge("bottom", bottom); _edge("right", right)

def set_cell_margins(cell, top=120, bottom=120, left=120, right=120):
    _, _, _, qn, OxmlElement, _, _, _ = _docx_primitives()
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcMar = tcPr.find(qn("w:tcMar"))
    if tcMar is None:
        tcMar = OxmlElement("w:tcMar")
        tcPr.append(tcMar)
    for side, val in (("top", top), ("bottom", bottom), ("start", left), ("end", right)):
        elt = tcMar.find(qn(f"w:{side}"))
        if elt is None:
            elt = OxmlElement(f"w:{side}")
            tcMar.append(elt)
        elt.set(qn("w:w"), str(val))
        elt.set(qn("w:type"), "dxa")

def add_picture_paragraph(doc, img_bytes: bytes, width_mm: Optional[int] = 140):
    _, _, _, _, _, Mm, _, _ = _docx_primitives()
    par = doc.add_paragraph()
    stream = BytesIO(img_bytes)
    run = par.add_run()
    if width_mm:
        run.add_picture(stream, width=Mm(width_mm))
    else:
        run.add_picture(stream)
    return par

def add_image_boxed(doc, img_bytes: bytes, width_mm: Optional[int] = 140,
                    title: Optional[str] = None, caption: Optional[str] = None):
    _, _, _, _, _, Mm, _, _ = _docx_primitives()
    table = doc.add_table(rows=1, cols=1)
    table.autofit = True
    cell = table.cell(0, 0)
    set_cell_borders(cell, top=("single", 12, "000000"),
                           left=("single", 12, "000000"),
                           bottom=("single", 12, "000000"),
                           right=("single", 12, "000000"))
    set_cell_margins(cell, top=120, bottom=120, left=160, right=160)
    if title:
        pt = cell.add_paragraph()
        rt = add_ko_run(pt, _strip_controls(_strip_html(title)))
        rt.bold = True
    pimg = cell.add_paragraph()
    run = pimg.add_run()
    stream = BytesIO(img_bytes)
    if width_mm:
        run.add_picture(stream, width=Mm(width_mm))
    else:
        run.add_picture(stream)
    if caption:
        pc = cell.add_paragraph()
        add_ko_run(pc, _strip_controls(_strip_html(caption)))
    doc.add_paragraph("")

def render_chart_png(chart: ChartData) -> bytes:
    if not chart.labels or not chart.datasets:
        raise ValueError("chart_data is empty")
    plt = _plt()
    fig, ax = plt.subplots(figsize=(6, 3.8), dpi=150)
    x = range(len(chart.labels))
    if chart.type == "bar":
        n = len(chart.datasets); total_width = 0.8
        bar_w = total_width / max(n, 1)
        offs = [(-total_width/2 + bar_w/2) + i*bar_w for i in range(n)]
        for i, ds in enumerate(chart.datasets):
            y = ds.data if ds and ds.data else [0]*len(chart.labels)
            ax.bar([xx+offs[i] for xx in x], y, width=bar_w, label=(ds.label or f"S{i+1}"))
        ax.set_xticks(list(x)); ax.set_xticklabels(chart.labels, rotation=0)
    elif chart.type == "line":
        for i, ds in enumerate(chart.datasets):
            y = ds.data if ds and ds.data else [0]*len(chart.labels)
            ax.plot(chart.labels, y, marker="o", label=(ds.label or f"S{i+1}"))
    else:
        n = len(chart.datasets); total_width = 0.8
        bar_w = total_width / max(n, 1)
        offs = [(-total_width/2 + bar_w/2) + i*bar_w for i in range(n)]
        for i, ds in enumerate(chart.datasets):
            y = ds.data if ds and ds.data else [0]*len(chart.labels)
            ax.bar([xx+offs[i] for xx in x], y, width=bar_w, label=(ds.label or f"S{i+1}"))
        ax.set_xticks(list(x)); ax.set_xticklabels(chart.labels, rotation=0)
    if chart.title: ax.set_title(chart.title)
    ax.grid(True, axis="y", linestyle="--", alpha=0.3)
    ax.legend(loc="best")
    fig.tight_layout()
    buf = BytesIO(); fig.savefig(buf, format="png"); plt.close(fig); buf.seek(0)
    return buf.getvalue()

def add_table_boxed(doc, headers: List[str], rows: List[List[Any]], title: Optional[str] = None):
    _, _, _, _, _, _, WD_TABLE_ALIGNMENT, WD_ALIGN_PARAGRAPH = _docx_primitives()
    outer = doc.add_table(rows=1, cols=1); outer.autofit = True
    ocell = outer.cell(0, 0)
    set_cell_borders(ocell, top=("single", 12, "000000"),
                           left=("single", 12, "000000"),
                           bottom=("single", 12, "000000"),
                           right=("single", 12, "000000"))
    set_cell_margins(ocell, top=120, bottom=120, left=160, right=160)
    if title:
        pt = ocell.add_paragraph(); rt = add_ko_run(pt, _strip_controls(_strip_html(title))); rt.bold = True
    ncols = len(headers) if headers else (max((len(r) for r in rows), default=0))
    tbl = ocell.add_table(rows=1, cols=max(ncols, 1)); tbl.alignment = WD_TABLE_ALIGNMENT.CENTER; tbl.autofit = True
    if headers:
        hdr = tbl.rows[0]
        for j, h in enumerate(headers):
            ph = hdr.cells[j].paragraphs[0]; rh = add_ko_run(ph, _strip_controls(_strip_html(str(h)))); rh.bold = True
            ph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for r in rows:
        row = tbl.add_row()
        for j in range(len(row.cells)):
            val = r[j] if j < len(r) else ""
            p = row.cells[j].paragraphs[0]
            add_ko_run(p, _strip_controls(_strip_html(str(val))))
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph("")

# ── 주어진 문장 값 취득(스키마 불일치 안전) ───────────────────────────
def _get_given_sentence(obj) -> Optional[str]:
    # 1) snake_case
    val = getattr(obj, "given_sentence", None)
    if isinstance(val, str):
        return val
    # 2) camelCase
    val = getattr(obj, "givenSentence", None)
    if isinstance(val, str):
        return val
    # 3) dict-like 백업
    try:
        d = obj.__dict__
        v = d.get("given_sentence") or d.get("givenSentence")
        if isinstance(v, str):
            return v
    except Exception:
        pass
    # 4) pydantic v2 model_dump()
    try:
        d = obj.model_dump()
        v = d.get("given_sentence") or d.get("givenSentence")
        if isinstance(v, str):
            return v
    except Exception:
        pass
    return None

# ── 주어진 문장 박스 렌더러 ─────────────────────────────────────────
def add_given_sentence_boxed(doc, sent: Optional[str], title: Optional[str] = "[주어진 문장]"):
    if not (isinstance(sent, str) and sent.strip() != ""):
        return
    from docx.shared import Mm  # not used but ok to import here if extended later
    table = doc.add_table(rows=1, cols=1); table.autofit = True
    cell = table.cell(0, 0)
    set_cell_borders(cell, top=("single", 12, "000000"),
                           left=("single", 12, "000000"),
                           bottom=("single", 12, "000000"),
                           right=("single", 12, "000000"))
    set_cell_margins(cell, top=120, bottom=120, left=160, right=160)
    if title:
        p_title = cell.add_paragraph(); r_t = add_ko_run(p_title, _strip_controls(_strip_html(title))); r_t.bold = True
    p_body = cell.add_paragraph(); add_rich_ko(p_body, sent)
    doc.add_paragraph("")

def add_summary_template(doc, template: Optional[str]):
    if not template:
        return
    raw = _strip_controls(_strip_html_except_u(template))
    if not raw.strip():
        return
    for line in raw.split("\n"):
        if not line.strip():
            continue
        par = doc.add_paragraph(); add_rich_ko(par, line)

def add_summary_template_boxed(doc, template: Optional[str], title: Optional[str] = None):
    if not template:
        return
    raw = _strip_controls(_strip_html_except_u(template))
    if not raw.strip():
        return
    table = doc.add_table(rows=1, cols=1); table.autofit = True
    cell = table.cell(0, 0)
    set_cell_borders(cell, top=("single", 12, "000000"),
                           left=("single", 12, "000000"),
                           bottom=("single", 12, "000000"),
                           right=("single", 12, "000000"))
    set_cell_margins(cell, top=120, bottom=120, left=160, right=160)
    if title:
        p_title = cell.add_paragraph(); run_t = add_ko_run(p_title, _strip_controls(_strip_html(title))); run_t.bold = True
    for line in raw.split("\n"):
        if not line.strip():
            continue
        par = cell.add_paragraph(); add_rich_ko(par, line)
    doc.add_paragraph("")

def add_options(doc, options: List[str], optionsLabeled: Optional[List[LabeledOption]] = None):
    if optionsLabeled:
        for opt in optionsLabeled:
            par = doc.add_paragraph()
            add_symbol_run(par, (opt.label or "").strip()); add_ko_run(par, NBSP); add_rich_ko(par, opt.text or "")
    else:
        for i, raw in enumerate(options):
            par = doc.add_paragraph()
            add_symbol_run(par, circled_label(i)); add_ko_run(par, NBSP); add_rich_ko(par, raw)

def add_passage(doc, passage: Optional[str], passage_paragraphs: Optional[List[str]]):
    if passage_paragraphs and len(passage_paragraphs) > 0:
        for ptxt in passage_paragraphs:
            ptxt = _strip_controls(ptxt or "")
            if not ptxt.strip():
                continue
            par = doc.add_paragraph(); add_rich_ko(par, ptxt)
    elif passage:
        text = _strip_controls(_strip_html_except_u(passage))
        parts = [t for t in re.split(r"\n\s*\n", text) if t.strip()]
        for ptxt in parts:
            par = doc.add_paragraph(); add_rich_ko(par, ptxt)

def render_chart_or_table_into_doc(doc, maybe):
    if not maybe:
        return
    try:
        # --- 새로 추가: dict → 모델 인스턴스로 캐스팅 ---
        if isinstance(maybe, dict):
            if "headers" in maybe and "rows" in maybe:
                maybe = TableData(**maybe)
            else:
                maybe = ChartData(**maybe)
        # -------------------------------------------------

        if isinstance(maybe, TableData):
            add_table_boxed(doc, maybe.headers, maybe.rows, title=(maybe.title or "Table"))
        else:
            img_bytes = render_chart_png(maybe)  # ChartData
            add_image_boxed(doc, img_bytes, width_mm=150, title="Chart", caption=(maybe.title or None))
    except Exception:
        pass

def render_images_into_doc(doc, image_base64: Optional[str], images: Optional[List[ImageSpec]]):
    if image_base64:
        try:
            img_bytes = _data_url_to_bytes(image_base64)
            add_picture_paragraph(doc, img_bytes, width_mm=140)
        except Exception:
            pass
    if images:
        for im in images:
            try:
                img_bytes = _data_url_to_bytes(im.data_url)
                if im.boxed:
                    add_image_boxed(doc, img_bytes, width_mm=im.width_mm, title=im.title, caption=im.caption)
                else:
                    add_picture_paragraph(doc, img_bytes, width_mm=im.width_mm)
                    if im.caption:
                        pcap = doc.add_paragraph(); add_ko_run(pcap, _strip_controls(_strip_html(im.caption)))
            except Exception:
                continue

def render_subitem(doc, si: SubItem, idx: int, mode: str, parent_order: int):
    if getattr(si, "question", None):
        par_q = doc.add_paragraph()
        prefix = f"문항 {parent_order} - {idx} "
        run_prefix = add_ko_run(par_q, prefix); run_prefix.bold = True
        add_rich_ko(par_q, si.question)

    # 주어진 문장(스키마 가변 대응)
    sent = _get_given_sentence(si)
    if isinstance(sent, str) and sent.strip() != "":
        add_given_sentence_boxed(doc, sent, title="[주어진 문장]")

    if getattr(si, "summary_template", None):
        add_summary_template_boxed(doc, si.summary_template, title="Summary Template")
    render_chart_or_table_into_doc(doc, getattr(si, "chart_data", None))
    render_images_into_doc(doc, getattr(si, "image_base64", None), getattr(si, "images", None))

    if getattr(si, "optionsLabeled", None) or getattr(si, "options", None):
        add_options(doc, getattr(si, "options", []), getattr(si, "optionsLabeled", None))

    if mode in ("answer", "explain"):
        if getattr(si, "answer", None) not in (None, ""):
            p_ans = doc.add_paragraph(); add_ko_run(p_ans, f"정답: {str(si.answer).strip()}")
        if mode == "explain" and getattr(si, "explain", None):
            p_exp = doc.add_paragraph(); add_ko_run(p_exp, "해설: "); add_ko_run(p_exp, _strip_controls(_strip_html(si.explain)))
    doc.add_paragraph("")

# ── 진입점 ────────────────────────────────────────────────────────
def generate_docx(payload: ExportPayload) -> tuple[str, str]:
    from docx import Document
    from docx.oxml.ns import qn
    from docx.shared import Pt, RGBColor

    title = payload.title or "시험지"

    doc = Document()
    normal = doc.styles["Normal"]
    n_rPr = normal.element.get_or_add_rPr()
    n_rFonts = n_rPr.get_or_add_rFonts()
    normal.font.name = FONT_KO
    n_rFonts.set(qn("w:ascii"),    FONT_KO)
    n_rFonts.set(qn("w:hAnsi"),    FONT_KO)
    n_rFonts.set(qn("w:eastAsia"), FONT_KO)
    n_rFonts.set(qn("w:cs"),       FONT_KO)

    # 제목/설명
    doc.add_heading(_strip_controls(_strip_html(title)), level=0)
    if getattr(payload, "description", None):
        par = doc.add_paragraph(); add_ko_run(par, _strip_controls(_strip_html(payload.description)))
    doc.add_paragraph("")

    # 뒤로 모을 정답/해설
    answer_rows: List[Tuple[str, str]] = []
    explain_blocks: List[Tuple[str, str]] = []

    show_answers_in_body = (payload.mode in ("answer", "explain")) and (not getattr(payload, "answers_at_end", False))
    show_explain_in_body = (payload.mode == "explain") and (not getattr(payload, "explain_at_end", False))

    # 본문
    for it in sorted(payload.items, key=lambda x: x.order):
        # === 문항 헤더 (문항 N + item_name 같은 줄에 출력) ===
        from docx.shared import Pt, RGBColor

        p_head = doc.add_paragraph()
        # 왼쪽: 문항 번호 (기존 스타일)
        r_main = add_ko_run(p_head, f"문항 {it.order}")
        r_main.bold = True
        r_main.font.size = Pt(14)

        # 오른쪽: item_name (작고 파란색)
        item_name_raw = getattr(it, "item_name", None)
        item_name_clean = _strip_controls(_strip_html(item_name_raw)) if item_name_raw else ""
        if item_name_clean:
            add_ko_run(p_head, " ")  # 간격 하나
            r_sub = add_ko_run(p_head, f"문항유형: {item_name_clean}")
            r_sub.font.size = Pt(10)
            r_sub.font.color.rgb = RGBColor(0x6A, 0xA7, 0xDE)  # 흐린 파란색
        # ===================================================

        if getattr(it, "question", None):
            par = doc.add_paragraph(); add_rich_ko(par, it.question)

        add_passage(doc, getattr(it, "passage", None), getattr(it, "passage_paragraphs", None))

        # 아이템 레벨 주어진 문장(스키마 가변 대응)
        sent = _get_given_sentence(it)
        if isinstance(sent, str) and sent.strip() != "":
            add_given_sentence_boxed(doc, sent, title="[주어진 문장]")

        if getattr(it, "summary_template", None):
            add_summary_template_boxed(doc, it.summary_template, title="Summary Template")

        cd = getattr(it, "chart_data", None)
        if not getattr(it, "images", None) and not getattr(it, "image_base64", None) and cd:
            try:
                if isinstance(cd, TableData): add_table_boxed(doc, cd.headers, cd.rows, title=(cd.title or "Table"))
                else:
                    img_bytes = render_chart_png(cd)
                    add_image_boxed(doc, img_bytes, width_mm=150, title="Chart", caption=(cd.title or None))
            except Exception:
                pass

        if getattr(it, "image_base64", None):
            try:
                img_bytes = _data_url_to_bytes(it.image_base64); add_picture_paragraph(doc, img_bytes, width_mm=140)
            except Exception:
                pass

        if getattr(it, "subItems", None):
            subitem_body_mode = "student" if (getattr(payload, "answers_at_end", False) or getattr(payload, "explain_at_end", False)) else payload.mode
            for idx, si in enumerate(it.subItems, 1):
                render_subitem(doc, si, idx, mode=subitem_body_mode, parent_order=it.order)
                qid = f"{it.order}-{idx}"
                if getattr(payload, "answers_at_end", False) and getattr(si, "answer", None) not in (None, ""):
                    answer_rows.append((qid, str(si.answer).strip()))
                if getattr(payload, "explain_at_end", False) and getattr(si, "explain", None):
                    explain_blocks.append((qid, _strip_controls(_strip_html(si.explain))))
            doc.add_paragraph("")
            continue

        if getattr(it, "optionsLabeled", None) or getattr(it, "options", None):
            add_options(doc, getattr(it, "options", []), getattr(it, "optionsLabeled", None))

        if show_answers_in_body and getattr(it, "answer", None) not in (None, ""):
            par = doc.add_paragraph(); add_ko_run(par, f"정답: {str(it.answer).strip()}")
        if show_explain_in_body and getattr(it, "explain", None):
            par = doc.add_paragraph(); add_ko_run(par, "해설: "); add_ko_run(par, _strip_controls(_strip_html(it.explain)))

        if getattr(payload, "answers_at_end", False) and getattr(it, "answer", None) not in (None, ""):
            answer_rows.append((str(it.order), str(it.answer).strip()))
        if getattr(payload, "explain_at_end", False) and getattr(it, "explain", None):
            explain_blocks.append((str(it.order), _strip_controls(_strip_html(it.explain))))
        doc.add_paragraph("")

    # Appendix A. 정답표
    if getattr(payload, "answers_at_end", False) and answer_rows:
        doc.add_page_break(); doc.add_heading("Appendix A. 정답표", level=1)
        def _key_row(row: Tuple[str, str]):
            key = row[0]; parts = key.split("-")
            main = int(parts[0]) if parts[0].isdigit() else 0
            sub = int(parts[1]) if len(parts) > 1 and parts[1].isdigit() else 0
            return (main, sub)
        answer_rows_sorted = sorted(answer_rows, key=_key_row)
        add_table_boxed(doc, ["문항", "정답"], [[qid, ans] for qid, ans in answer_rows_sorted], title=None)

    # Appendix B. 해설
    if getattr(payload, "explain_at_end", False) and explain_blocks:
        doc.add_heading("Appendix B. 해설", level=1)
        def _key_block(row: Tuple[str, str]):
            key = row[0]; parts = key.split("-")
            main = int(parts[0]) if parts[0].isdigit() else 0
            sub = int(parts[1]) if len(parts) > 1 and parts[1].isdigit() else 0
            return (main, sub)
        explain_blocks_sorted = sorted(explain_blocks, key=_key_block)
        last_main = None
        for qid, exp in explain_blocks_sorted:
            main = qid.split("-")[0]
            if last_main != main:
                doc.add_heading(f"문항 {main}", level=2); last_main = main
            if "-" in qid:
                p_head = doc.add_paragraph(); run = add_ko_run(p_head, f"[{qid}] "); run.bold = True
            p = doc.add_paragraph(); add_ko_run(p, exp); doc.add_paragraph("")

    # 파일 저장
    tmp = NamedTemporaryFile(delete=False, suffix=".docx")
    tmp_path = tmp.name; tmp.close(); doc.save(tmp_path)
    return tmp_path, f"{(payload.title or '시험지')}.docx"
