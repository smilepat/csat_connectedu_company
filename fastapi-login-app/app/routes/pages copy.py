# app/routes/pages.py
from fastapi import APIRouter, Request, HTTPException, Header, Depends
from pydantic import BaseModel, Field
from typing import List, Optional, Dict, Any, get_args

import time
import json
import redis

from app.core.logging import logger, log_action
from app.services.http_client import post_java  # async

pages_router = APIRouter(prefix="/api/pages", tags=["pages"])

# ✅ Redis: 문자열로 받도록
r = redis.Redis(host="localhost", port=6379, db=0, decode_responses=True)

# ✅ 인증/세션
def token_required(authorization: Optional[str] = Header(None)):  # ← Header 주입
    if not authorization or not authorization.startswith("Bearer "):
        raise HTTPException(status_code=401, detail="인증 토큰이 필요합니다")
    token = authorization.replace("Bearer ", "")
    user_data = r.get(f"auth:{token}")
    if not user_data:
        raise HTTPException(status_code=401, detail="토큰이 만료되었거나 유효하지 않습니다")
    try:
        # decode_responses=True 이므로 바로 loads 가능
        return json.loads(user_data)
    except Exception:
        raise HTTPException(status_code=401, detail="유효하지 않은 토큰 데이터입니다")

# ===== 요청 모델 동일 =====
class ComposeAdd(BaseModel):
    question_seq: int
    display_order: int
    section_label: Optional[str] = None
    points: Optional[int] = None
    note: Optional[str] = None

class ComposeReorder(BaseModel):
    question_seq: int
    display_order: int

class ComposePayload(BaseModel):
    user_seq: int
    page_id: int
    add: List[ComposeAdd] = Field(default_factory=list)
    reorder: List[ComposeReorder] = Field(default_factory=list)
    remove: List[int] = Field(default_factory=list)

# ===== 공통 프록시 동일(로깅 포함) =====
async def _proxy(
    request: Request,
    java_path: str,
    payload: Dict[str, Any],
    action: str,
    user: Dict[str, Any],
):
    payload = dict(payload or {})
    payload["user_seq"] = user.get("user_seq")

    t0 = time.time()
    try:
        res = await post_java(java_path, payload, request=request)
    except Exception as e:
        elapsed = int((time.time() - t0) * 1000)
        log_action(
            logger,
            getattr(request.state, "req_id", None),
            payload.get("user_seq"),
            payload.get("page_id"),
            action,
            elapsed,
            "9",
            f"UPSTREAM_CALL_FAILED: {e}",
        )
        raise HTTPException(status_code=502, detail="Java API 요청 실패")

    elapsed = int((time.time() - t0) * 1000)
    log_action(
        logger,
        getattr(request.state, "req_id", None),
        payload.get("user_seq"),
        payload.get("page_id"),
        action,
        elapsed,
        res.get("result"),
        None if res.get("result") != "9" else res.get("message"),
    )

    if "result" not in res:
        return {"result": "9", "message": "INVALID_UPSTREAM_PAYLOAD"}
    return res

# ===== 엔드포인트 동일 =====
@pages_router.post("/list")
async def list_pages(request: Request, payload: dict, user=Depends(token_required)):
    return await _proxy(request, "/pages/list", payload, "list", user)

@pages_router.post("/detail")
async def detail_page(request: Request, payload: dict, user=Depends(token_required)):
    return await _proxy(request, "/pages/detail", payload, "detail", user)

@pages_router.post("/add")
async def add_page(request: Request, payload: dict, user=Depends(token_required)):
    return await _proxy(request, "/pages/add", payload, "add", user)

@pages_router.post("/edit")
async def edit_page(request: Request, payload: dict, user=Depends(token_required)):
    return await _proxy(request, "/pages/edit", payload, "edit", user)

@pages_router.post("/delete")
async def delete_page(request: Request, payload: dict, user=Depends(token_required)):
    return await _proxy(request, "/pages/delete", payload, "delete", user)

@pages_router.post("/compose")
async def compose_page(request: Request, payload: ComposePayload, user=Depends(token_required)):
    data = payload.model_dump()
    return await _proxy(request, "/pages/compose", data, "compose", user)

# --- 추가 import ---
from fastapi.responses import FileResponse
from fastapi import BackgroundTasks, Request, Depends, HTTPException  # ✅ Request/Depends/HTTPException 사용중
from pydantic import BaseModel, Field
from typing import List, Optional, Any, Dict, Union, Tuple
from docx import Document
from docx.text.paragraph import Paragraph
from docx.text.run import Run
from docx.oxml.ns import qn  # eastAsia 폰트 지정을 위해 필요
from docx.oxml import OxmlElement
from docx.shared import Mm
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
import base64
import matplotlib
matplotlib.use("Agg")  # 헤드리스 환경
import matplotlib.pyplot as plt
from io import BytesIO
from tempfile import NamedTemporaryFile
import os
import re
import time

# (필요시 주석 해제/정리)
from app.core.logging import logger, log_action  # ✅ 코드에서 사용중

# --- 한글/심볼 폰트 상수 ---
FONT_KO = "Malgun Gothic"
FONT_SYMBOL = "Segoe UI Symbol"

# --- 간단 HTML 제거 ---
def _strip_html(s: Optional[str]) -> str:
    if not s:
        return ""
    s = re.sub(r"(?i)<br\s*/?>", "\n", s)
    s = re.sub(r"<[^>]+>", "", s)
    return re.sub(r"\n{3,}", "\n\n", s).strip()

def _strip_controls(s: str) -> str:
    # ZWSP 등 제어문자 제거
    return re.sub(r"[\u200B-\u200D\uFEFF]", "", s)

# --- ①~⑳ 라벨 유니코드 생성 ---
def circled_label(i: int) -> str:
    base = 0x2460
    return chr(base + i) if 0 <= i <= 19 else f"{i+1}."

def _data_url_to_bytes(data_url: str) -> bytes:
    # "data:image/png;base64,....."
    m = re.match(r"^data:(image/\w+);base64,(.+)$", data_url)
    if not m:
        # header가 빠진 순수 base64만 왔을 수도 있음
        if re.match(r"^[A-Za-z0-9+/=\s]+$", data_url or ""):
            return base64.b64decode(data_url)
        raise ValueError("Invalid data URL")
    b64 = m.group(2)
    return base64.b64decode(b64)

def add_picture_paragraph(doc: Document, img_bytes: bytes, width_mm: Optional[int] = 140):
    par = doc.add_paragraph()
    stream = BytesIO(img_bytes)
    run = par.add_run()
    if width_mm:
        run.add_picture(stream, width=Mm(width_mm))
    else:
        run.add_picture(stream)
    return par

def add_image_boxed(doc: Document, img_bytes: bytes, width_mm: Optional[int] = 140,
                    title: Optional[str] = None, caption: Optional[str] = None):
    # 1x1 테이블로 박스 생성
    table = doc.add_table(rows=1, cols=1)
    table.autofit = True
    cell = table.cell(0, 0)

    # 테두리
    set_cell_borders(cell, top=("single", 12, "000000"),
                           left=("single", 12, "000000"),
                           bottom=("single", 12, "000000"),
                           right=("single", 12, "000000"))
    set_cell_margins(cell, top=120, bottom=120, left=160, right=160)

    if title:
        pt = cell.add_paragraph()
        rt = add_ko_run(pt, _strip_controls(_strip_html(title)))
        rt.bold = True

    # 이미지
    pimg = cell.add_paragraph()
    run = pimg.add_run()
    stream = BytesIO(img_bytes)
    if width_mm:
        run.add_picture(stream, width=Mm(width_mm))
    else:
        run.add_picture(stream)

    # 캡션
    if caption:
        pc = cell.add_paragraph()
        add_ko_run(pc, _strip_controls(_strip_html(caption)))

    # 표 이후 공간
    doc.add_paragraph("")

# --- Run 헬퍼: KO/심볼 각각 전용 (rPr/rFonts 안전 생성) ✅ ---
def add_ko_run(par: Paragraph, text: str) -> Run:
    run = par.add_run(text)  # ✅ 오타 수정: add_run
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.get_or_add_rFonts()
    run.font.name = FONT_KO
    rFonts.set(qn("w:ascii"),    FONT_KO)
    rFonts.set(qn("w:hAnsi"),    FONT_KO)
    rFonts.set(qn("w:eastAsia"), FONT_KO)
    rFonts.set(qn("w:cs"),       FONT_KO)
    return run

def add_symbol_run(par: Paragraph, text: str) -> Run:
    run = par.add_run(text)  # ✅ 오타 수정: add_run
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.get_or_add_rFonts()
    run.font.name = FONT_SYMBOL
    # 라벨 Run은 eastAsia까지 심볼 폰트로 고정 ✅
    rFonts.set(qn("w:ascii"),    FONT_SYMBOL)
    rFonts.set(qn("w:hAnsi"),    FONT_SYMBOL)
    rFonts.set(qn("w:eastAsia"), FONT_SYMBOL)
    rFonts.set(qn("w:cs"),       FONT_SYMBOL)
    return run

def set_cell_borders(cell, top=("single", 12, "000000"),
                          left=("single", 12, "000000"),
                          bottom=("single", 12, "000000"),
                          right=("single", 12, "000000")):
    """
    cell 테두리를 설정합니다.
    각 인자: (val, size, color)  e.g., ("single", 12, "000000")
    size 단위: eights of a point (1pt = 8)
    """
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = tcPr.find(qn("w:tcBorders"))
    if tcBorders is None:
        tcBorders = OxmlElement("w:tcBorders")
        tcPr.append(tcBorders)

    def _edge(tag, spec):
        if spec is None:
            return
        val, sz, color = spec
        edge = tcBorders.find(qn(f"w:{tag}"))
        if edge is None:
            edge = OxmlElement(f"w:{tag}")
            tcBorders.append(edge)
        edge.set(qn("w:val"), val)
        edge.set(qn("w:sz"), str(sz))
        edge.set(qn("w:color"), color)

    _edge("top", top)
    _edge("left", left)
    _edge("bottom", bottom)
    _edge("right", right)

def set_cell_margins(cell, top=120, bottom=120, left=120, right=120):
    """
    셀 내부 여백 설정 (twips, 1pt = 20 twips). 기본 6pt 정도.
    """
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcMar = tcPr.find(qn("w:tcMar"))
    if tcMar is None:
        tcMar = OxmlElement("w:tcMar")
        tcPr.append(tcMar)

    for side, val in (("top", top), ("bottom", bottom), ("start", left), ("end", right)):
        elt = tcMar.find(qn(f"w:{side}"))
        if elt is None:
            elt = OxmlElement(f"w:{side}")
            tcMar.append(elt)
        elt.set(qn("w:w"), str(val))
        elt.set(qn("w:type"), "dxa")

class ImageSpec(BaseModel):
    data_url: str                     # "data:image/png;base64,..." 형태
    caption: Optional[str] = None     # 캡션(선택)
    width_mm: Optional[int] = 140     # 폭(mm) 기본 140mm
    boxed: Optional[bool] = False     # 테두리 박스 여부
    title: Optional[str] = None       # 박스 상단 제목(선택, 예: "Chart")
    
class ChartDataset(BaseModel):
    label: Optional[str] = None
    data: List[float]

class ChartData(BaseModel):
    type: str  # 'bar' | 'line'
    title: Optional[str] = None
    labels: List[str]
    datasets: List[ChartDataset]

class TableData(BaseModel):
    headers: List[str]
    rows: List[List[Any]]
    title: Optional[str] = None

# --- 내보내기용 모델 ---
class LabeledOption(BaseModel):
    label: str
    text: str

class SubItem(BaseModel):
    question: Optional[str] = None
    options: List[str] = Field(default_factory=list)
    optionsLabeled: Optional[List[LabeledOption]] = None
    answer: Optional[Any] = None
    explain: Optional[str] = None
    given_sentence: Optional[str] = None
    summary_template: Optional[str] = None
    chart_data: Optional[Union[ChartData, TableData]] = None
    image_base64: Optional[str] = None
    images: Optional[List[ImageSpec]] = None

class ExportItem(BaseModel):
    order: int
    question: Optional[str] = None
    passage: Optional[str] = None
    passage_paragraphs: Optional[List[str]] = None
    options: List[str] = Field(default_factory=list)
    optionsLabeled: Optional[List[LabeledOption]] = None
    answer: Optional[Any] = None  # 숫자/문자 모두 허용
    explain: Optional[str] = None
    meta: Optional[Dict[str, Any]] = None
    chart_data: Optional[Union[ChartData, TableData]] = None
    image_base64: Optional[str] = None
    images: Optional[List[ImageSpec]] = None
    # (앞서 진행한 summary_template도 쓰고 있다면 여기에 포함)
    summary_template: Optional[str] = None    
    subItems: Optional[List[SubItem]] = None

class ExportPayload(BaseModel):
    title: str = "시험지"
    description: Optional[str] = None
    mode: str = Field(pattern="^(student|answer|explain)$")
    items: List[ExportItem] = Field(default_factory=list)
    # ✅ 추가: 본문에는 정답/해설을 숨기고 뒤에 모아서 출력
    answers_at_end: bool = True
    explain_at_end: bool = True

def render_chart_png(chart: ChartData) -> bytes:
    # 안전 처리: 비어있으면 빈 이미지 대신 예외
    if not chart.labels or not chart.datasets:
        raise ValueError("chart_data is empty")

    fig, ax = plt.subplots(figsize=(6, 3.8), dpi=150)  # 적당한 비율
    x = range(len(chart.labels))

    if chart.type == "bar":
        # 막대 간 그룹형 배치
        n = len(chart.datasets)
        total_width = 0.8
        bar_w = total_width / max(n, 1)
        offs = [(-total_width/2 + bar_w/2) + i*bar_w for i in range(n)]
        for i, ds in enumerate(chart.datasets):
            y = ds.data if ds and ds.data else [0]*len(chart.labels)
            ax.bar([xx+offs[i] for xx in x], y, width=bar_w, label=(ds.label or f"S{i+1}"))
    elif chart.type == "line":
        for i, ds in enumerate(chart.datasets):
            y = ds.data if ds and ds.data else [0]*len(chart.labels)
            ax.plot(chart.labels, y, marker="o", label=(ds.label or f"S{i+1}"))
    else:
        # 기본은 bar로
        n = len(chart.datasets)
        total_width = 0.8
        bar_w = total_width / max(n, 1)
        offs = [(-total_width/2 + bar_w/2) + i*bar_w for i in range(n)]
        for i, ds in enumerate(chart.datasets):
            y = ds.data if ds and ds.data else [0]*len(chart.labels)
            ax.bar([xx+offs[i] for xx in x], y, width=bar_w, label=(ds.label or f"S{i+1}"))
        ax.set_xticks(list(x))
        ax.set_xticklabels(chart.labels, rotation=0)

    if chart.type == "bar":
        ax.set_xticks(list(x))
        ax.set_xticklabels(chart.labels, rotation=0)

    if chart.title:
        ax.set_title(chart.title)

    ax.grid(True, axis="y", linestyle="--", alpha=0.3)
    ax.legend(loc="best")
    fig.tight_layout()

    buf = BytesIO()
    fig.savefig(buf, format="png")
    plt.close(fig)
    buf.seek(0)
    return buf.getvalue()

def add_table_boxed(doc: Document, headers: List[str], rows: List[List[Any]], title: Optional[str] = None):
    # 1x1 테이블(박스) 생성 후 내부에 또 표를 넣는 방식으로 테두리/마진 일관 유지
    outer = doc.add_table(rows=1, cols=1)
    outer.autofit = True
    ocell = outer.cell(0, 0)
    set_cell_borders(ocell, top=("single", 12, "000000"),
                           left=("single", 12, "000000"),
                           bottom=("single", 12, "000000"),
                           right=("single", 12, "000000"))
    set_cell_margins(ocell, top=120, bottom=120, left=160, right=160)

    if title:
        pt = ocell.add_paragraph()
        rt = add_ko_run(pt, _strip_controls(_strip_html(title)))
        rt.bold = True

    # 실제 데이터 표
    ncols = len(headers) if headers else (max((len(r) for r in rows), default=0))
    tbl = ocell.add_table(rows=1, cols=max(ncols, 1))
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    tbl.autofit = True

    # 헤더
    if headers:
        hdr = tbl.rows[0]
        for j, h in enumerate(headers):
            ph = hdr.cells[j].paragraphs[0]
            rh = add_ko_run(ph, _strip_controls(_strip_html(str(h))))
            rh.bold = True
            ph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    else:
        # 헤더 없으면 빈 헤더 한 줄
        pass

    # 바디
    for r in rows:
        row = tbl.add_row()
        for j in range(len(row.cells)):
            val = r[j] if j < len(r) else ""
            p = row.cells[j].paragraphs[0]
            add_ko_run(p, _strip_controls(_strip_html(str(val))))
            # 가운데/왼쪽 정렬은 필요 시 조절
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # 박스 아래 여백
    doc.add_paragraph("")

def add_given_sentence(doc: Document, sent: Optional[str]):
    if not sent:
        return
    par = doc.add_paragraph()
    add_ko_run(par, _strip_controls(_strip_html(f"[주어진 문장] {sent}")))

def render_chart_or_table_into_doc(doc: Document, maybe):
    if not maybe:
        return
    try:
        if isinstance(maybe, TableData):
            add_table_boxed(doc, maybe.headers, maybe.rows, title=(maybe.title or "Table"))
        else:
            img_bytes = render_chart_png(maybe)  # ChartData
            add_image_boxed(doc, img_bytes, width_mm=150, title="Chart", caption=(maybe.title or None))
    except Exception:
        pass

def render_images_into_doc(doc: Document, image_base64: Optional[str], images: Optional[List[ImageSpec]]):
    # 단일 base64
    if image_base64:
        try:
            img_bytes = _data_url_to_bytes(image_base64)
            add_picture_paragraph(doc, img_bytes, width_mm=140)
        except Exception:
            pass
    # 여러 ImageSpec
    if images:
        for im in images:
            try:
                img_bytes = _data_url_to_bytes(im.data_url)
                if im.boxed:
                    add_image_boxed(doc, img_bytes, width_mm=im.width_mm, title=im.title, caption=im.caption)
                else:
                    add_picture_paragraph(doc, img_bytes, width_mm=im.width_mm)
                    if im.caption:
                        pcap = doc.add_paragraph()
                        add_ko_run(pcap, _strip_controls(_strip_html(im.caption)))
            except Exception:
                continue

def render_subitem(doc: Document, si: "SubItem", idx: int, mode: str, parent_order: int):
    # 질문 머리말: "문항 {parent_order} - {idx} "
    if si.question:
        par_q = doc.add_paragraph()
        # 머리말(굵게) + 공백 후 질문 본문
        prefix = f"문항 {parent_order} - {idx} "
        run_prefix = add_ko_run(par_q, prefix)
        run_prefix.bold = True
        add_ko_run(par_q, _strip_controls(_strip_html(si.question)))
        
    # 주어진 문장 / 요약 템플릿 / 자료(표·차트·이미지)
    add_given_sentence(doc, si.given_sentence)
    if si.summary_template:
        add_summary_template_boxed(doc, si.summary_template, title="Summary Template")
    render_chart_or_table_into_doc(doc, si.chart_data)
    render_images_into_doc(doc, si.image_base64, si.images)

    # 보기
    if si.optionsLabeled or si.options:
        add_options(doc, si.options, si.optionsLabeled)

    # 정답/해설 (본문 출력은 mode에 따름)
    if mode in ("answer", "explain"):
        if si.answer is not None and str(si.answer).strip():
            p_ans = doc.add_paragraph()
            add_ko_run(p_ans, f"정답: {str(si.answer).strip()}")
        if mode == "explain" and si.explain:
            p_exp = doc.add_paragraph()
            add_ko_run(p_exp, "해설: ")
            add_ko_run(p_exp, _strip_controls(_strip_html(si.explain)))

    # 부문항 간 여백
    doc.add_paragraph("")

# --- 문단 추가(문단 배열/문자열 모두 허용) ---
def add_passage(doc: Document, passage: Optional[str], passage_paragraphs: Optional[List[str]]):
    if passage_paragraphs and len(passage_paragraphs) > 0:
        for p in passage_paragraphs:
            p = _strip_controls(_strip_html(p))
            if not p:
                continue
            par = doc.add_paragraph()
            add_ko_run(par, p)
    elif passage:
        text = _strip_controls(_strip_html(passage))
        # \n\n 기준으로 문단 분리
        parts = [t for t in re.split(r"\n\s*\n", text) if t.strip()]
        for p in parts:
            par = doc.add_paragraph()
            add_ko_run(par, p)

# --- 보기 추가(라벨 분리: 심볼 폰트 + 한글 폰트) ---
NBSP = "\u00A0"

def _strip_html_except_u(s: Optional[str]) -> str:
    """<u>만 남기고 나머지 태그는 제거. <br> -> 줄바꿈."""
    if not s:
        return ""
    s = re.sub(r"(?i)<br\s*/?>", "\n", s)
    # u 태그만 임시 토큰으로 보존
    s = re.sub(r"(?i)</u>", "__CLOSE_U__", s)
    s = re.sub(r"(?i)<u>", "__OPEN_U__", s)
    # 나머지 태그 제거
    s = re.sub(r"<[^>]+>", "", s)
    # u 태그 복원
    s = s.replace("__OPEN_U__", "<u>").replace("__CLOSE_U__", "</u>")
    return s

def _preserve_spaces(text: str) -> str:
    """연속 공백을 NBSP로 보존 (DOCX에서 일반 공백은 연속시 붙습니다)."""
    # 탭은 공백 4개 정도로
    text = text.replace("\t", "    ")
    # 연속 공백을 NBSP로 교체
    return re.sub(r" {2,}", lambda m: NBSP * len(m.group(0)), text)

def add_summary_template(doc: Document, template: Optional[str]):
    """summary_template을 한 단락으로 출력. <u>…</u>는 밑줄 처리."""
    if not template:
        return
    raw = _strip_controls(_strip_html_except_u(template))
    if not raw.strip():
        return

    # 여러 줄(줄바꿈)도 지원: 줄마다 별도 paragraph
    lines = raw.split("\n")
    for line in lines:
        if not line.strip():
            continue
        par = doc.add_paragraph()
        # <u>…</u> 구간을 기준으로 분리
        parts = re.split(r"(?i)(<u>.*?</u>)", line)
        for part in parts:
            if not part:
                continue
            m = re.match(r"(?i)^<u>(.*?)</u>$", part)
            if m:
                txt = _preserve_spaces(m.group(1))
                run = add_ko_run(par, txt)
                run.underline = True
            else:
                txt = _preserve_spaces(part)
                add_ko_run(par, txt)

def add_summary_template_boxed(doc: Document, template: Optional[str], title: Optional[str] = None):
    """
    summary_template을 1x1 테이블 셀 안에 넣고 테두리를 표시합니다.
    <u>…</u> 밑줄 유지, 여러 줄 지원.
    """
    if not template:
        return
    raw = _strip_controls(_strip_html_except_u(template))
    if not raw.strip():
        return

    # 표 생성 (1x1)
    table = doc.add_table(rows=1, cols=1)
    table.autofit = True
    cell = table.cell(0, 0)

    # 테두리 & 내부 여백
    set_cell_borders(cell, top=("single", 12, "000000"),
                           left=("single", 12, "000000"),
                           bottom=("single", 12, "000000"),
                           right=("single", 12, "000000"))
    set_cell_margins(cell, top=120, bottom=120, left=160, right=160)

    # 제목(선택): "Summary Template" 같은 헤더
    if title:
        p_title = cell.add_paragraph()
        run_t = add_ko_run(p_title, _strip_controls(_strip_html(title)))
        run_t.bold = True

    # 줄 단위로 파싱하여 밑줄 유지
    lines = raw.split("\n")
    for line in lines:
        if not line.strip():
            continue
        par = cell.add_paragraph()
        parts = re.split(r"(?i)(<u>.*?</u>)", line)
        for part in parts:
            if not part:
                continue
            m = re.match(r"(?i)^<u>(.*?)</u>$", part)
            if m:
                txt = _preserve_spaces(m.group(1))
                run = add_ko_run(par, txt)
                run.underline = True
            else:
                txt = _preserve_spaces(part)
                add_ko_run(par, txt)

    # 표 아래 여백용 빈 문단(선택)
    doc.add_paragraph("")

def add_options(doc: Document, options: List[str], optionsLabeled: Optional[List[LabeledOption]] = None):
    if optionsLabeled:
        for opt in optionsLabeled:
            par = doc.add_paragraph()
            add_symbol_run(par, (opt.label or "").strip())                         # ① (심볼 Run)
            add_ko_run(par, NBSP + _strip_controls(_strip_html(opt.text or "")))   # 본문 (KO Run)
    else:
        for i, raw in enumerate(options):
            par = doc.add_paragraph()
            add_symbol_run(par, circled_label(i))                                  # ① (심볼 Run)
            add_ko_run(par, NBSP + _strip_controls(_strip_html(raw)))              # 본문 (KO Run)

# --- DOCX 생성 엔드포인트 ---
@pages_router.post("/export_docx")
async def export_docx(
    payload: ExportPayload,
    request: Request,
    background: BackgroundTasks,
    user=Depends(token_required)
):
    """
    프론트에서 현재 화면의 모드/문항을 그대로 보내면 .docx로 반환합니다.
    POST /api/pages/export_docx
    body: ExportPayload
    """
    title = payload.title or "시험지"
    t0 = time.time()
    try:
        doc = Document()

        # 문서 기본 글꼴/동아시아 폰트 강제 (안전 생성) ✅
        normal = doc.styles["Normal"]
        n_rPr = normal.element.get_or_add_rPr()
        n_rFonts = n_rPr.get_or_add_rFonts()
        normal.font.name = FONT_KO
        n_rFonts.set(qn("w:ascii"),    FONT_KO)
        n_rFonts.set(qn("w:hAnsi"),    FONT_KO)
        n_rFonts.set(qn("w:eastAsia"), FONT_KO)
        n_rFonts.set(qn("w:cs"),       FONT_KO)

        # 제목/설명
        doc.add_heading(_strip_controls(_strip_html(title)), level=0)
        if payload.description:
            par = doc.add_paragraph()
            add_ko_run(par, _strip_controls(_strip_html(payload.description)))
        doc.add_paragraph("")  # spacing

        # ✅ 정답/해설 누적 버퍼
        answer_rows: List[Tuple[str, str]] = []   # [("1", "④"), ("1-1","B") ...]
        explain_blocks: List[Tuple[str, str]] = [] # [("1", "해설 본문..."), ("1-1", "...")]

        # 본문에서 정답/해설 즉시 출력 여부(기존 모드 유지 조건)
        show_answers_in_body = (payload.mode in ("answer", "explain")) and (not payload.answers_at_end)
        show_explain_in_body = (payload.mode == "explain") and (not payload.explain_at_end)

        # 본문 문항
        for it in sorted(payload.items, key=lambda x: x.order):
            # 문항 머리말
            doc.add_heading(f"문항 {it.order}", level=1)

            # 질문(있으면)
            if it.question:
                par = doc.add_paragraph()
                add_ko_run(par, _strip_controls(_strip_html(it.question)))

            # 지문(문단 배열 우선)
            add_passage(doc, it.passage, it.passage_paragraphs)

            # ⬇️ 요약 템플릿 (있을 때만)
            if it.summary_template:
                add_summary_template_boxed(doc, it.summary_template, title="Summary Template")

            # ⬇️ 차트/표/이미지 렌더
            cd = getattr(it, "chart_data", None)

            # ⬇️ 차트 렌더 & 삽입 (images가 없고 chart_data가 있으면 서버에서 생성)
            if not getattr(it, "images", None) and not getattr(it, "image_base64", None) and cd:
                try:
                    if isinstance(cd, TableData):
                        add_table_boxed(doc, cd.headers, cd.rows, title=(cd.title or "Table"))
                    else:
                        # ChartData로 간주
                        img_bytes = render_chart_png(cd)
                        add_image_boxed(doc, img_bytes, width_mm=150,
                                        title="Chart", caption=(cd.title or None))
                except Exception:
                    pass  # 실패해도 문항 진행

            # 1) 단일 image_base64 (빠른 경로)
            if getattr(it, "image_base64", None):
                try:
                    img_bytes = _data_url_to_bytes(it.image_base64)
                    # 박스 없이 곧바로
                    add_picture_paragraph(doc, img_bytes, width_mm=140)
                except Exception:
                    pass  # 실패해도 문항 진행은 계속

            # ✅ 셋트 문항 처리
            if it.subItems:
                # 본문에는 항상 '문제만' 출력할지 여부 결정
                # answers_at_end 또는 explain_at_end 중 하나라도 True면 본문은 student 모드로만 렌더
                subitem_body_mode = "student" if (payload.answers_at_end or payload.explain_at_end) else payload.mode

                for idx, si in enumerate(it.subItems, 1):
                    # 본문 렌더
                    render_subitem(doc, si, idx, mode=subitem_body_mode, parent_order=it.order)

                    # ✅ 뒤로 모을 데이터 수집
                    qid = f"{it.order}-{idx}"

                    # 정답 수집
                    if payload.answers_at_end and si.answer not in (None, ""):
                        answer_rows.append((qid, str(si.answer).strip()))

                    # 해설 수집
                    if payload.explain_at_end and si.explain:
                        explain_blocks.append((qid, _strip_controls(_strip_html(si.explain))))

                doc.add_paragraph("")
                continue

            # 단문항: 보기 출력
            if it.optionsLabeled or it.options:
                add_options(doc, it.options, it.optionsLabeled)

            # ✅ 정답/해설 (본문 즉시 출력 or 뒤로 수집)
            # 1) 본문 즉시 출력 (기존 호환)
            if show_answers_in_body and it.answer not in (None, ""):
                par = doc.add_paragraph()
                add_ko_run(par, f"정답: {str(it.answer).strip()}")
            if show_explain_in_body and it.explain:
                par = doc.add_paragraph()
                add_ko_run(par, "해설: ")
                add_ko_run(par, _strip_controls(_strip_html(it.explain)))

            # 2) 뒤로 수집
            if payload.answers_at_end and it.answer not in (None, ""):
                answer_rows.append((str(it.order), str(it.answer).strip()))
            if payload.explain_at_end and it.explain:
                explain_blocks.append((str(it.order), _strip_controls(_strip_html(it.explain))))

            doc.add_paragraph("")  # 문항 간 여백

        # ====== Appendix A: 정답표 ======
        if payload.answers_at_end and answer_rows:
            doc.add_page_break()
            doc.add_heading("Appendix A. 정답표", level=1)

            # 정렬: '1', '1-1', '1-2' 순서
            def _key_row(row: Tuple[str, str]):
                key = row[0]
                parts = key.split("-")
                main = int(parts[0]) if parts[0].isdigit() else 0
                sub = int(parts[1]) if len(parts) > 1 and parts[1].isdigit() else 0
                return (main, sub)

            answer_rows_sorted = sorted(answer_rows, key=_key_row)
            headers = ["문항", "정답"]   # 필요하면 ["문항","정답","배점","비고"]
            rows = [[qid, ans] for qid, ans in answer_rows_sorted]
            add_table_boxed(doc, headers, rows, title=None)

        # ====== Appendix B: 해설 ======
        if payload.explain_at_end and explain_blocks:
            doc.add_heading("Appendix B. 해설", level=1)

            def _key_block(row: Tuple[str, str]):
                key = row[0]
                parts = key.split("-")
                main = int(parts[0]) if parts[0].isdigit() else 0
                sub = int(parts[1]) if len(parts) > 1 and parts[1].isdigit() else 0
                return (main, sub)

            explain_blocks_sorted = sorted(explain_blocks, key=_key_block)
            last_main = None
            for qid, exp in explain_blocks_sorted:
                main = qid.split("-")[0]
                if last_main != main:
                    doc.add_heading(f"문항 {main}", level=2)
                    last_main = main
                # 부문항이면 하위 제목 표시
                if "-" in qid:
                    # 소제목 라벨
                    p_head = doc.add_paragraph()
                    run = add_ko_run(p_head, f"[{qid}] ")
                    run.bold = True
                p = doc.add_paragraph()
                add_ko_run(p, exp)
                doc.add_paragraph("")

        # 임시 파일 저장 + 응답
        tmp = NamedTemporaryFile(delete=False, suffix=".docx")
        tmp_path = tmp.name
        tmp.close()
        doc.save(tmp_path)

        # 로깅
        elapsed = int((time.time() - t0) * 1000)
        log_action(
            logger,
            getattr(request.state, "req_id", None),
            user.get("user_seq"),
            None,
            "export_docx",
            elapsed,
            "0",
            None,
        )

        # 응답 후 파일 삭제
        background.add_task(lambda p: os.remove(p) if os.path.exists(p) else None, tmp_path)

        return FileResponse(
            tmp_path,
            media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            filename=f"{title}.docx",
        )

    except Exception as e:
        elapsed = int((time.time() - t0) * 1000)
        log_action(
            logger,
            getattr(request.state, "req_id", None),
            user.get("user_seq"),
            None,
            "export_docx",
            elapsed,
            "9",
            f"EXPORT_FAILED: {e}",
        )
        raise HTTPException(status_code=500, detail="DOCX 생성 실패")
